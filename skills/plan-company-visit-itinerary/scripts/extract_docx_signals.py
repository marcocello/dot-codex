#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


def color_value(color) -> str | None:
    if color is None or color.rgb is None:
        return None
    return str(color.rgb)


def run_record(run) -> dict:
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": bool(run.underline) if run.underline is not None else None,
        "fontColor": color_value(run.font.color),
        "highlight": str(run.font.highlight_color) if run.font.highlight_color is not None else None,
    }


def paragraph_record(paragraph, index: int, location: str) -> dict:
    return {
        "index": index,
        "location": location,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style is not None else None,
        "runs": [run_record(run) for run in paragraph.runs if run.text],
    }


def extract(path: Path) -> dict:
    document = Document(path)
    paragraphs = [
        paragraph_record(paragraph, index, f"paragraph:{index}")
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    ]
    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for cell_index, cell in enumerate(row.cells):
                cells.append({
                    "index": cell_index,
                    "text": cell.text,
                    "paragraphs": [
                        paragraph_record(paragraph, paragraph_index, f"table:{table_index}/row:{row_index}/cell:{cell_index}/paragraph:{paragraph_index}")
                        for paragraph_index, paragraph in enumerate(cell.paragraphs)
                    ],
                })
            rows.append({"index": row_index, "cells": cells})
        tables.append({"index": table_index, "rows": rows})
    return {
        "source": str(path),
        "paragraphCount": len(paragraphs),
        "tableCount": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract DOCX text and formatting signals for company-list normalization.")
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    payload = extract(args.input)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({"output": str(args.output), "paragraphs": payload["paragraphCount"], "tables": payload["tableCount"]}))


if __name__ == "__main__":
    main()
